from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from typing import Dict, Any


class DOCXGenerator:
    """Generate professional DOCX resumes"""
    
    def __init__(self, template: str = "modern", theme_color: str = "#3B82F6"):
        self.template = template
        self.theme_color = theme_color
        self.doc = Document()
        self.buffer = BytesIO()
    
    def hex_to_rgb(self, hex_color: str):
        """Convert hex color to RGB tuple"""
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    
    def add_heading(self, text: str, level: int = 1):
        """Add a styled heading"""
        heading = self.doc.add_heading(text, level=level)
        if level == 1:
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.runs[0]
        run.font.color.rgb = RGBColor(*self.hex_to_rgb(self.theme_color))
        return heading
    
    def add_paragraph(self, text: str, bold: bool = False, italic: bool = False):
        """Add a paragraph with optional styling"""
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(11)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        return para
    
    def generate(self, resume_data: Dict[str, Any]) -> BytesIO:
        """Generate DOCX resume"""
        
        # Personal Info
        personal_info = resume_data.get('personal_info', {})
        if personal_info.get('name'):
            name_para = self.doc.add_heading(personal_info['name'], level=1)
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact information
        contact_parts = []
        if personal_info.get('email'):
            contact_parts.append(personal_info['email'])
        if personal_info.get('phone'):
            contact_parts.append(personal_info['phone'])
        if personal_info.get('location'):
            contact_parts.append(personal_info['location'])
        
        if contact_parts:
            contact_para = self.doc.add_paragraph(' | '.join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Links
        links = []
        if personal_info.get('linkedin'):
            links.append(f"LinkedIn: {personal_info['linkedin']}")
        if personal_info.get('github'):
            links.append(f"GitHub: {personal_info['github']}")
        if personal_info.get('portfolio'):
            links.append(f"Portfolio: {personal_info['portfolio']}")
        
        if links:
            links_para = self.doc.add_paragraph(' | '.join(links))
            links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()  # Spacer
        
        # Summary/Objective
        if resume_data.get('summary'):
            self.add_heading('CAREER OBJECTIVE / PROFESSIONAL SUMMARY', level=2)
            self.add_paragraph(resume_data['summary'])
            self.doc.add_paragraph()
        elif resume_data.get('objective'):
            self.add_heading('OBJECTIVE', level=2)
            self.add_paragraph(resume_data['objective'])
            self.doc.add_paragraph()
        
        # Skills - MOVED UP
        if resume_data.get('skills'):
            self.add_heading('TECHNICAL SKILLS', level=2)
            skills = resume_data['skills']
            
            # Handle both formats: List[str] or List[{category, items}]
            if skills and isinstance(skills[0], dict):
                # New format with categories
                for skill_cat in skills:
                    category = skill_cat.get('category', '')
                    items = skill_cat.get('items', [])
                    if category and items:
                        skills_text = f"{category}: {', '.join(items)}"
                        self.add_paragraph(skills_text)
            else:
                # Old format: simple list
                skills_text = ', '.join(skills)
                self.add_paragraph(skills_text)
            self.doc.add_paragraph()
        
        # Projects - MOVED UP
        if resume_data.get('projects'):
            self.add_heading('PROJECTS', level=2)
            for proj in resume_data['projects']:
                self.add_paragraph(proj.get('name', 'Project'), bold=True)
                
                if proj.get('description'):
                    self.add_paragraph(proj['description'])
                
                if proj.get('technologies'):
                    tech_text = f"Technologies: {', '.join(proj['technologies'])}"
                    self.add_paragraph(tech_text, italic=True)
                
                if proj.get('url'):
                    self.add_paragraph(f"URL: {proj['url']}")
                
                self.doc.add_paragraph()
        
        # Experience / Internship Experience
        if resume_data.get('experience'):
            self.add_heading('INTERNSHIP / WORK EXPERIENCE', level=2)
            for exp in resume_data['experience']:
                # Position and company
                self.add_paragraph(
                    f"{exp.get('position', 'Position')} at {exp.get('company', 'Company')}",
                    bold=True
                )
                
                # Dates and location
                date_info = []
                if exp.get('start_date'):
                    date_range = f"{exp['start_date']} - {exp.get('end_date', 'Present') if not exp.get('current') else 'Present'}"
                    date_info.append(date_range)
                if exp.get('location'):
                    date_info.append(exp['location'])
                
                if date_info:
                    self.add_paragraph(' | '.join(date_info), italic=True)
                
                # Description (can be string or list)
                description = exp.get('description')
                if description:
                    if isinstance(description, list):
                        for item in description:
                            if item:  # Skip empty strings
                                para = self.doc.add_paragraph(item, style='List Bullet')
                    else:
                        self.add_paragraph(description)
                
                # Achievements
                if exp.get('achievements'):
                    for achievement in exp['achievements']:
                        para = self.doc.add_paragraph(achievement, style='List Bullet')
                
                self.doc.add_paragraph()  # Spacer
        
        # Certifications
        if resume_data.get('certifications'):
            self.add_heading('CERTIFICATIONS', level=2)
            for cert in resume_data['certifications']:
                cert_text = f"{cert.get('name', 'Certification')} - {cert.get('issuer', 'Issuer')}"
                if cert.get('date'):
                    cert_text += f" ({cert['date']})"
                self.add_paragraph(cert_text)
            self.doc.add_paragraph()
        
        # Education - MOVED TO END
        if resume_data.get('education'):
            self.add_heading('EDUCATION', level=2)
            for edu in resume_data['education']:
                # Degree
                degree_text = f"{edu.get('degree', 'Degree')}"
                field = edu.get('field_of_study') or edu.get('field', '')
                if field:
                    degree_text += f" in {field}"
                self.add_paragraph(degree_text, bold=True)
                
                # Institution
                self.add_paragraph(edu.get('institution', 'Institution'))
                
                # Dates and grade
                date_info = []
                if edu.get('graduation_date'):
                    date_info.append(edu['graduation_date'])
                elif edu.get('start_date'):
                    date_range = f"{edu['start_date']} - {edu.get('end_date', 'Present')}"
                    date_info.append(date_range)
                if edu.get('gpa') or edu.get('grade'):
                    gpa = edu.get('gpa') or edu.get('grade')
                    date_info.append(f"GPA: {gpa}")
                
                if date_info:
                    self.add_paragraph(' | '.join(date_info), italic=True)
                
                if edu.get('description'):
                    self.add_paragraph(edu['description'])
                
                self.doc.add_paragraph()
        
        # Languages
        if resume_data.get('languages'):
            self.add_heading('LANGUAGES', level=2)
            languages_text = ', '.join(resume_data['languages'])
            self.add_paragraph(languages_text)
            self.doc.add_paragraph()
        
        # Awards
        if resume_data.get('awards'):
            self.add_heading('AWARDS & HONORS', level=2)
            for award in resume_data['awards']:
                para = self.doc.add_paragraph(award, style='List Bullet')
        
        # Save to buffer
        self.doc.save(self.buffer)
        self.buffer.seek(0)
        return self.buffer


def generate_docx_resume(resume_data: Dict[str, Any], template: str = "modern", theme_color: str = "#3B82F6") -> BytesIO:
    """Main function to generate DOCX resume"""
    generator = DOCXGenerator(template, theme_color)
    return generator.generate(resume_data)
